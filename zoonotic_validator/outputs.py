# outputs.py --> GENERACIÓN DE FICHEROS DE SALIDA PARA EL PIPELINE DE VALIDACIÓN DE DATOS ZOONÓTICOS
# Este módulo contiene funciones para generar los archivos de salida del pipeline de validación de datos zoonóticos, incluyendo:
    # - Un Excel con los errores detectados.
    # - Una copia del Excel original con las celdas que contienen errores resaltadas.
    # - Un informe en Word que resume los errores encontrados y proporciona detalles para su revisión.

from collections import Counter # Importamos Counter para contar el número de incidencias por tipo de error, lo que nos permite generar un resumen claro en el informe de Word.
from pathlib import Path # Importamos Path para manejar las rutas de los archivos de manera más robusta y compatible con diferentes sistemas operativos.

import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook, Workbook
from openpyxl.styles import PatternFill


ERROR_FILL = PatternFill(fill_type="solid", fgColor="FF6B6B") # Un color de relleno rojo claro para resaltar las celdas con errores en el Excel marcado. Se puede personalizar según las preferencias de visualización.

# Esta función toma un DataFrame con los errores detectados y lo guarda en un archivo Excel. El parámetro index=False se utiliza para evitar que se guarde una columna adicional con los índices del DataFrame, lo que hace que el archivo de salida sea más limpio y fácil de leer. El nombre del archivo de salida se especifica a través del parámetro output_file.
def save_errors_to_excel(errors_df: pd.DataFrame, output_file: str) -> None:
    """Save the error table as a simple Excel file."""
    errors_df.to_excel(output_file, index=False)
# ============================================================

# Esta función crea una copia del archivo Excel original y resalta las celdas que contienen errores a nivel de celda. Carga el libro de Excel, accede a la hoja especificada, y luego itera sobre los errores que son a nivel de celda para aplicar un relleno de color a las celdas correspondientes. Finalmente, guarda el libro modificado con un nuevo nombre especificado por output_file. Es importante destacar que solo se pueden resaltar los errores que están asociados a celdas específicas, ya que los errores estructurales como la falta de columnas no tienen una celda concreta que resaltar.
def create_marked_excel(
    input_file: str,
    sheet_name,
    errors_df: pd.DataFrame,
    output_file: str,
) -> None:
    """Create a copy of the original workbook and highlight invalid cells.

    Supports both .xls and .xlsx file formats.
    Only cell-level errors can be highlighted because structural errors
    such as missing columns do not point to a specific existing cell.
    """
    # Detectar si es archivo .xls o .xlsx
    is_xls = input_file.lower().endswith('.xls')
    
    if is_xls:
        # Para archivos .xls: leer con pandas (usa xlrd automáticamente) y recrear con openpyxl
        df_data = pd.read_excel(input_file, sheet_name=sheet_name)
        
        # Crear nuevo workbook
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = df_data.index.name or "Sheet1"
        
        # Escribir encabezados
        for col_idx, col_name in enumerate(df_data.columns, start=1):
            worksheet.cell(row=1, column=col_idx, value=col_name)
        
        # Escribir datos
        for row_idx, row in enumerate(df_data.itertuples(index=False), start=2):
            for col_idx, value in enumerate(row, start=1):
                worksheet.cell(row=row_idx, column=col_idx, value=value)
    else:
        # Para archivos .xlsx: usar openpyxl directamente
        workbook = load_workbook(input_file)
        worksheet = workbook[workbook.sheetnames[sheet_name] if isinstance(sheet_name, int) else sheet_name]

    # Aplicar resaltado a los errores a nivel de celda
    if not errors_df.empty:
        cell_level_errors = errors_df[errors_df["is_cell_level"] == True]
        for _, error in cell_level_errors.iterrows():
            if pd.notna(error.get("excel_column")) and pd.notna(error.get("excel_row")):
                worksheet.cell(
                    row=int(error["excel_row"]),
                    column=int(error["excel_column"]),
                ).fill = ERROR_FILL

    workbook.save(output_file)
# ============================================================

# Esta función genera un informe en formato Word que incluye un resumen de los errores detectados, un desglose por tipo de error, y una tabla detallada con cada incidencia. 
# Utiliza la biblioteca python-docx para crear el documento, agregar encabezados, párrafos y tablas. 
# El informe se guarda con el nombre especificado en output_file. 
# Este informe es útil para que los usuarios puedan revisar fácilmente los errores encontrados y entender qué aspectos necesitan ser corregidos en el Excel original.
def create_word_report(errors_df: pd.DataFrame, output_file: str, input_file: str = None) -> None:
    """Generate a clear Word report with summary and detailed tables."""
    # Diccionario con descripciones de códigos de error
    error_descriptions = {
        "E001": "Columna obligatoria faltante",
        "E002": "Campo obligatorio vacío",
        "E003": "Valor debe ser exactamente el especificado",
        "E004": "Año en repYear no coincide con el año de Mapping_Options",
        "E005": "Valor prohibido detectado",
        "E006": "Formato numérico inválido",
        "E007": "Formato recId inválido: CCAA en MAYÚSCULAS, agente válido, números según fila",
        "E008": "Fila completa duplicada",
    }
    
    # Diccionario con mensajes específicos de revisión para cada tipo de error
    error_review_messages = {
        "E001": "Agregue la columna faltante al fichero Excel.",
        "E002": "Rellene el campo obligatorio que está vacío.",
        "E003": "Corrija el valor para que coincida exactamente con el especificado.",
        "E004": "Corrija el año en repYear para que coincida con el año de la hoja Mapping_Options (celda A1).",
        "E005": "No puede ser \"unspecified\", \"desconocido\" ni \"unknown\".",
        "E006": "Corrija el formato del número (verifique decimales, separadores y tipo de dato).",
        "E007": "Corrija el recId siguiendo el patrón CCAA_AGENTE_###. Verifique: (1) CCAA en MAYÚSCULAS, (2) Agente zoon. válido, (3) Números según fila del Excel.",
        "E008": "Elimine la fila duplicada o verifique que los datos sean correctamente únicos.",
    }
    
    document = Document()
    document.add_heading("Informe de validación del Excel", level=1)
    
    # Agregar nombre del archivo si está disponible
    if input_file:
        file_name = Path(input_file).name
        p = document.add_paragraph(f"Archivo validado: {file_name}")
        # Aplicar formato en negrita al párrafo
        for run in p.runs:
            run.bold = True

    if errors_df.empty:
        document.add_paragraph("No se han detectado errores en las validaciones generales.")
        document.save(output_file)
        return

    document.add_paragraph(
        "Se han detectado incidencias en la revisión general del fichero. "
        "A continuación se muestra un resumen y el detalle de lo que conviene revisar."
    )

    document.add_heading("Resumen", level=2)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Concepto"
    header_cells[1].text = "Valor"

    row = summary_table.add_row().cells
    row[0].text = "Número total de incidencias"
    row[1].text = str(len(errors_df))

    structural_count = int((errors_df["is_cell_level"] == False).sum())
    cell_count = int((errors_df["is_cell_level"] == True).sum())

    row = summary_table.add_row().cells
    row[0].text = "Incidencias estructurales"
    row[1].text = str(structural_count)

    row = summary_table.add_row().cells
    row[0].text = "Incidencias en celdas concretas"
    row[1].text = str(cell_count)

    document.add_heading("Resumen por tipo de error", level=2)
    grouped = Counter(errors_df["error_code"])
    grouped_table = document.add_table(rows=1, cols=3)
    grouped_table.style = "Table Grid"
    grouped_header = grouped_table.rows[0].cells
    grouped_header[0].text = "Código"
    grouped_header[1].text = "Descripción"
    grouped_header[2].text = "Número de casos"

    for code, count in sorted(grouped.items()):
        row = grouped_table.add_row().cells
        row[0].text = str(code)
        row[1].text = error_descriptions.get(str(code), "Error desconocido")
        row[2].text = str(count)

    document.add_heading("Detalle de incidencias", level=2)
    detail_table = document.add_table(rows=1, cols=6)
    detail_table.style = "Table Grid"
    detail_header = detail_table.rows[0].cells
    detail_header[0].text = "Fila Excel"
    detail_header[1].text = "Campo"
    detail_header[2].text = "Valor detectado"
    detail_header[3].text = "Código"
    detail_header[4].text = "Incidencia"
    detail_header[5].text = "Qué revisar"

    for _, error in errors_df.iterrows():
        row = detail_table.add_row().cells
        row[0].text = str(error["excel_row"])
        row[1].text = str(error["field"])
        row[2].text = "null" if pd.isna(error["value"]) else str(error["value"])
        row[3].text = str(error["error_code"])
        row[4].text = str(error["message"])
        error_code = str(error["error_code"])
        row[5].text = error_review_messages.get(error_code, "Revise y corrija el error según lo indicado.")

    document.save(output_file)
